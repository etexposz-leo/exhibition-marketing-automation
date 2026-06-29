"""Event Workspace API Routes."""

import os
import json
import shutil
from datetime import datetime
from pathlib import Path
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, Request, UploadFile, File, Form, Body
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.models.models import Event, EventDocument, EventDocumentChunk, EventQuery
from app.services.rag_service import rag_service

router = APIRouter()

# Upload directory - use local path
BASE_DIR = Path(__file__).parent.parent.parent
UPLOAD_DIR = BASE_DIR / "data" / "uploads" / "events"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def require_sms_verified(request: Request) -> int:
    """Require user to be authenticated and SMS verified. Returns user_id."""
    user_id = request.session.get("user_id")
    if not user_id:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if not request.session.get("sms_verified"):
        raise HTTPException(status_code=403, detail="Phone verification required")
    return user_id


# ==================== Event CRUD ====================


@router.get("/events/stats")
async def get_event_stats(request: Request, db: Session = Depends(get_db)):
    """Get event statistics for dashboard."""
    user_id = require_sms_verified(request)
    
    total_events = db.query(Event).filter(Event.user_id == user_id).count()
    
    # Active events (upcoming or ongoing)
    today = datetime.now().date()
    active_events = db.query(Event).filter(
        Event.user_id == user_id,
        Event.status == "active",
        (Event.start_date >= today) | (Event.end_date >= today)
    ).count()
    
    # Total documents
    total_docs = db.query(EventDocument).filter(
        EventDocument.user_id == user_id
    ).count()
    
    # Total queries
    total_queries = db.query(EventQuery).filter(
        EventQuery.user_id == user_id
    ).count()
    
    return {
        "total_events": total_events,
        "active_events": active_events,
        "total_documents": total_docs,
        "total_queries": total_queries
    }


@router.post("/events")
async def create_event(
    request: Request,
    event_name: str = Form(...),
    venue: Optional[str] = Form(None),
    city: Optional[str] = Form(None),
    country: Optional[str] = Form(None),
    start_date: Optional[str] = Form(None),
    end_date: Optional[str] = Form(None),
    organizer: Optional[str] = Form(None),
    website: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    """Create a new event."""
    user_id = require_sms_verified(request)
    
    from datetime import date
    
    # Parse dates
    start_date_obj = None
    end_date_obj = None
    if start_date:
        try:
            start_date_obj = datetime.strptime(start_date, "%Y-%m-%d").date()
        except ValueError:
            pass
    if end_date:
        try:
            end_date_obj = datetime.strptime(end_date, "%Y-%m-%d").date()
        except ValueError:
            pass
    
    event = Event(
        user_id=user_id,
        event_name=event_name,
        venue=venue,
        city=city,
        country=country,
        start_date=start_date_obj,
        end_date=end_date_obj,
        organizer=organizer,
        website=website,
        status="active"
    )
    db.add(event)
    db.commit()
    db.refresh(event)
    
    return {
        "id": event.id,
        "event_name": event.event_name,
        "venue": event.venue,
        "city": event.city,
        "country": event.country,
        "start_date": str(event.start_date) if event.start_date else None,
        "end_date": str(event.end_date) if event.end_date else None,
        "organizer": event.organizer,
        "website": event.website,
        "status": event.status,
        "created_at": event.created_at.isoformat()
    }


@router.get("/events")
async def list_events(request: Request, db: Session = Depends(get_db)):
    """List all events for the current user."""
    user_id = require_sms_verified(request)
    
    events = db.query(Event).filter(
        Event.user_id == user_id
    ).order_by(Event.created_at.desc()).all()
    
    return [
        {
            "id": e.id,
            "event_name": e.event_name,
            "venue": e.venue,
            "city": e.city,
            "country": e.country,
            "start_date": str(e.start_date) if e.start_date else None,
            "end_date": str(e.end_date) if e.end_date else None,
            "organizer": e.organizer,
            "website": e.website,
            "status": e.status,
            "created_at": e.created_at.isoformat()
        }
        for e in events
    ]


@router.get("/events/{event_id}")
async def get_event(event_id: int, request: Request, db: Session = Depends(get_db)):
    """Get a specific event."""
    user_id = require_sms_verified(request)
    
    event = db.query(Event).filter(
        Event.id == event_id,
        Event.user_id == user_id
    ).first()
    
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")
    
    # Get document count
    doc_count = db.query(EventDocument).filter(
        EventDocument.event_id == event_id
    ).count()
    
    return {
        "id": event.id,
        "event_name": event.event_name,
        "venue": event.venue,
        "city": event.city,
        "country": event.country,
        "start_date": str(event.start_date) if event.start_date else None,
        "end_date": str(event.end_date) if event.end_date else None,
        "organizer": event.organizer,
        "website": event.website,
        "status": event.status,
        "document_count": doc_count,
        "created_at": event.created_at.isoformat()
    }


@router.put("/events/{event_id}")
async def update_event(
    event_id: int,
    request: Request,
    event_name: str = Form(...),
    venue: Optional[str] = Form(None),
    city: Optional[str] = Form(None),
    country: Optional[str] = Form(None),
    start_date: Optional[str] = Form(None),
    end_date: Optional[str] = Form(None),
    organizer: Optional[str] = Form(None),
    website: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    """Update an event."""
    user_id = require_sms_verified(request)
    
    event = db.query(Event).filter(
        Event.id == event_id,
        Event.user_id == user_id
    ).first()
    
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")
    
    # Parse dates
    if start_date:
        try:
            event.start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
        except ValueError:
            pass
    if end_date:
        try:
            event.end_date = datetime.strptime(end_date, "%Y-%m-%d").date()
        except ValueError:
            pass
    
    event.event_name = event_name
    event.venue = venue
    event.city = city
    event.country = country
    event.organizer = organizer
    event.website = website
    
    db.commit()
    db.refresh(event)
    
    return {
        "id": event.id,
        "event_name": event.event_name,
        "message": "Event updated successfully"
    }


@router.delete("/events/{event_id}")
async def delete_event(event_id: int, request: Request, db: Session = Depends(get_db)):
    """Delete an event and its documents."""
    user_id = require_sms_verified(request)
    
    event = db.query(Event).filter(
        Event.id == event_id,
        Event.user_id == user_id
    ).first()
    
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")
    
    # Delete associated documents and chunks
    db.query(EventDocumentChunk).filter(
        EventDocumentChunk.event_id == event_id
    ).delete()
    
    db.query(EventDocument).filter(
        EventDocument.event_id == event_id
    ).delete()
    
    db.query(EventQuery).filter(
        EventQuery.event_id == event_id
    ).delete()
    
    # Delete event
    db.delete(event)
    db.commit()
    
    # Delete files
    event_dir = UPLOAD_DIR / str(event_id)
    if event_dir.exists():
        shutil.rmtree(event_dir)
    
    return {"message": "Event deleted successfully"}


# ==================== Document Management ====================


@router.post("/events/{event_id}/documents")
async def upload_document(
    event_id: int,
    request: Request,
    file: UploadFile = File(...),
    document_type: str = Form("exhibitor_manual"),
    db: Session = Depends(get_db)
):
    """Upload a document to an event."""
    user_id = require_sms_verified(request)
    
    # Verify event exists and belongs to user
    event = db.query(Event).filter(
        Event.id == event_id,
        Event.user_id == user_id
    ).first()
    
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")
    
    # Validate file type
    allowed_types = {"pdf", "docx", "txt"}
    file_ext = file.filename.split(".")[-1].lower() if "." in file.filename else ""
    
    if file_ext not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed. Allowed types: {', '.join(allowed_types)}"
        )
    
    # Create event directory
    event_dir = UPLOAD_DIR / str(event_id)
    event_dir.mkdir(parents=True, exist_ok=True)
    
    # Generate unique filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_filename = f"{timestamp}_{file.filename}"
    file_path = event_dir / safe_filename
    
    # Save file
    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Create document record
    document = EventDocument(
        event_id=event_id,
        user_id=user_id,
        filename=safe_filename,
        original_filename=file.filename,
        file_type=file_ext,
        file_size=len(content),
        file_path=str(file_path),
        document_type=document_type,
        status="processing"
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    
    # Process document in background (extract text and chunk)
    await process_document_async(document.id, file_path, file_ext, db)
    
    return {
        "id": document.id,
        "filename": document.original_filename,
        "status": document.status,
        "message": "Document uploaded successfully"
    }


async def process_document_async(document_id: int, file_path: Path, file_ext: str, db: Session):
    """Process document: extract text and create chunks."""
    try:
        # Extract text based on file type
        text = ""
        
        if file_ext == "txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        elif file_ext == "pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text())
                text = "\n\n".join(text_parts)
            except ImportError:
                text = "[PDF processing requires pypdf package]"
        
        elif file_ext == "docx":
            try:
                from docx import Document
                doc = Document(file_path)
                text = "\n\n".join([p.text for p in doc.paragraphs])
            except ImportError:
                text = "[DOCX processing requires python-docx package]"
        
        # Split into chunks
        chunk_size = 500
        chunks = []
        words = text.split()
        
        current_chunk = []
        current_length = 0
        
        for word in words:
            current_chunk.append(word)
            current_length += len(word) + 1
            
            if current_length >= chunk_size:
                chunk_text = " ".join(current_chunk)
                chunks.append(chunk_text)
                current_chunk = []
                current_length = 0
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        # Store chunks
        document = db.query(EventDocument).filter(EventDocument.id == document_id).first()
        if document:
            for i, chunk_text in enumerate(chunks):
                chunk = EventDocumentChunk(
                    event_id=document.event_id,
                    user_id=document.user_id,
                    document_id=document_id,
                    chunk_index=i,
                    content=chunk_text
                )
                db.add(chunk)
            
            document.chunk_count = len(chunks)
            document.status = "indexed"
            db.commit()
    
    except Exception as e:
        document = db.query(EventDocument).filter(EventDocument.id == document_id).first()
        if document:
            document.status = "error"
            db.commit()


@router.get("/events/{event_id}/documents")
async def list_documents(event_id: int, request: Request, db: Session = Depends(get_db)):
    """List all documents for an event."""
    user_id = require_sms_verified(request)
    
    # Verify event exists and belongs to user
    event = db.query(Event).filter(
        Event.id == event_id,
        Event.user_id == user_id
    ).first()
    
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")
    
    documents = db.query(EventDocument).filter(
        EventDocument.event_id == event_id
    ).order_by(EventDocument.created_at.desc()).all()
    
    return [
        {
            "id": d.id,
            "filename": d.original_filename,
            "file_type": d.file_type,
            "document_type": d.document_type,
            "status": d.status,
            "chunk_count": d.chunk_count,
            "file_size": d.file_size,
            "created_at": d.created_at.isoformat()
        }
        for d in documents
    ]


@router.delete("/events/{event_id}/documents/{document_id}")
async def delete_document(
    event_id: int,
    document_id: int,
    request: Request,
    db: Session = Depends(get_db)
):
    """Delete a document."""
    user_id = require_sms_verified(request)
    
    document = db.query(EventDocument).filter(
        EventDocument.id == document_id,
        EventDocument.event_id == event_id,
        EventDocument.user_id == user_id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file
    if document.file_path and Path(document.file_path).exists():
        Path(document.file_path).unlink()
    
    # Delete chunks
    db.query(EventDocumentChunk).filter(
        EventDocumentChunk.document_id == document_id
    ).delete()
    
    # Delete document
    db.delete(document)
    db.commit()
    
    return {"message": "Document deleted successfully"}


# ==================== Event AI Assistant ====================


@router.post("/events/{event_id}/query")
async def query_event(
    event_id: int,
    request: Request,
    body: dict = Body(...),
    db: Session = Depends(get_db)
):
    """Query the event knowledge base."""
    user_id = require_sms_verified(request)
    
    # Verify event exists and belongs to user
    event = db.query(Event).filter(
        Event.id == event_id,
        Event.user_id == user_id
    ).first()
    
    if not event:
        raise HTTPException(status_code=404, detail="Event not found")
    
    # Get all chunks for this event
    chunks = db.query(EventDocumentChunk).filter(
        EventDocumentChunk.event_id == event_id,
        EventDocumentChunk.user_id == user_id
    ).all()
    
    if not chunks:
        return {
            "answer": "No documents have been uploaded for this event yet. Please upload exhibitor manuals or venue rules to enable the AI assistant.",
            "sources": [],
            "provider": "mock"
        }
    
    # Get document info for chunks
    document_ids = set(c.document_id for c in chunks)
    documents = {
        d.id: d for d in db.query(EventDocument).filter(
            EventDocument.id.in_(document_ids)
        ).all()
    }
    
    # Format chunks for RAG
    chunk_dicts = [
        {
            "document_id": c.document_id,
            "document_title": documents.get(c.document_id, {}).original_filename or "Unknown",
            "chunk_index": c.chunk_index,
            "content": c.content
        }
        for c in chunks
    ]
    
    # Get params from body
    question = body.get("question", "")
    provider = body.get("provider", "mock")
    
    # Query RAG service
    result = await rag_service.query(question, chunk_dicts, provider)
    
    # Save query history
    query_record = EventQuery(
        event_id=event_id,
        user_id=user_id,
        question=question,
        answer=result["answer"],
        sources_json=json.dumps(result["sources"]),
        provider=result["provider"]
    )
    db.add(query_record)
    db.commit()
    
    return {
        "id": query_record.id,
        "question": question,
        "answer": result["answer"],
        "sources": result["sources"],
        "provider": result["provider"],
        "created_at": query_record.created_at.isoformat()
    }


@router.get("/events/{event_id}/queries")
async def get_event_queries(event_id: int, request: Request, db: Session = Depends(get_db)):
    """Get query history for an event."""
    user_id = require_sms_verified(request)
    
    queries = db.query(EventQuery).filter(
        EventQuery.event_id == event_id,
        EventQuery.user_id == user_id
    ).order_by(EventQuery.created_at.desc()).limit(50).all()
    
    return [
        {
            "id": q.id,
            "question": q.question,
            "answer": q.answer,
            "provider": q.provider,
            "created_at": q.created_at.isoformat()
        }
        for q in queries
    ]



